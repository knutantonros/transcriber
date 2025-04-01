# Python imports
import os
import streamlit as st
import hashlib
import tempfile
from PIL import Image

# External imports
from docx import Document
from pydub import AudioSegment

# Local imports
from utils.audio_utils import convert_to_mono_and_compress
from utils.transcribe import transcribe_with_kb_whisper
from utils.summarize import summarize_text_openai
import config as c

### INITIAL VARIABLES

# Skapar mappar om de inte existerar
# Använd absoluta sökvägar för att undvika problem i olika miljöer
base_dir = os.path.dirname(os.path.abspath(__file__))
audio_dir = os.path.join(base_dir, "audio")
text_dir = os.path.join(base_dir, "text")

os.makedirs(audio_dir, exist_ok=True)  # Där ljudfiler lagras för transkribering
os.makedirs(text_dir, exist_ok=True)   # Där transkriberade dokument lagras

# Kontrollera och sätt standardvärden om de inte finns i session_state
if "transcribed" not in st.session_state:  # Transkriptionens resultat
    st.session_state["transcribed"] = None
if "summarized" not in st.session_state:  # Sammanfattningens resultat
    st.session_state["summarized"] = None
if "transcribe_model" not in st.session_state:  # Vilken Whisper modell som ska användas
    st.session_state["transcribe_model"] = "KB Whisper Tiny"
if "file_name_converted" not in st.session_state:  # Ljudfilens namn
    st.session_state["file_name_converted"] = None

# Kontrollerar om uppladdad ljudfil har transkriberats
def compute_file_hash(uploaded_file):
    """Beräkna MD5-hash för en fil för att kontrollera om den har ändrats"""
    hasher = hashlib.md5()
    
    for chunk in iter(lambda: uploaded_file.read(4096), b""):
        hasher.update(chunk)
    uploaded_file.seek(0)  # Återställ filpekaren till början
    
    return hasher.hexdigest()

### HUVUDAPP ###########################

# Sidkonfiguration
st.set_page_config(
    page_title="Ljudtranskribering & Sammanfattning",
    page_icon="🎙️",
    layout="centered",
    initial_sidebar_state="auto"
)

# Inaktivera PyTorch-filövervakaren som orsakar problem
os.environ["STREAMLIT_GLOBAL_WATCHER_MAX_FILE_SIZE"] = "0"

def main():
    ### SIDOFÄLT
    
    st.sidebar.header("Inställningar")
    st.sidebar.markdown("")
    
    # Textfält för OpenAI API-nyckel
    api_key = st.sidebar.text_input(
        "Ange din OpenAI API-nyckel för sammanfattning",
        type="password",
        help="Krävs för sammanfattningsfunktionen. Din API-nyckel lagras endast i din aktiva session och skickas inte till någon server förutom OpenAI."
    )
    
    # Spara API-nyckeln i sessionen
    if api_key:
        st.session_state["openai_api_key"] = api_key
    
    # Dropdown-meny - välj Whisper-modell
    transcribe_model = st.sidebar.selectbox(
        "Välj transkriptionsmodell", 
        [
            "KB Whisper Tiny",
            "KB Whisper Base",
            "KB Whisper Small",
            "KB Whisper Medium", 
            "KB Whisper Large"
        ],
        index=[
            "KB Whisper Tiny",
            "KB Whisper Base",
            "KB Whisper Small",
            "KB Whisper Medium", 
            "KB Whisper Large"
        ].index(st.session_state["transcribe_model"]),
        help="Mindre modeller (Tiny, Base) är snabbare men mindre exakta. Större modeller (Medium, Large) är långsammare men mer exakta."
    )
    
    model_map_transcribe_model = {
        "KB Whisper Large": "kb-whisper-large",
        "KB Whisper Medium": "kb-whisper-medium",
        "KB Whisper Small": "kb-whisper-small",
        "KB Whisper Base": "kb-whisper-base",
        "KB Whisper Tiny": "kb-whisper-tiny"
    }
    
    # Val för sammanfattningslängd
    summary_length = st.sidebar.select_slider(
        "Sammanfattningslängd",
        options=["Mycket kort", "Kort", "Medium", "Lång", "Mycket lång"],
        value="Medium"
    )
    
    # Uppdatera session_state
    st.session_state["transcribe_model"] = transcribe_model
    
    st.sidebar.markdown("---")
    st.sidebar.markdown(f"Version: {c.app_version}")
    
    ### HUVUDSIDA
    
    # Titel
    st.title("Ljudtranskribering & Sammanfattning")
    st.markdown("### Ladda upp ljudfiler och få transkription + sammanfattning")
    
    st.markdown(f"""**Vald modell:** {st.session_state["transcribe_model"]}""")
    
    # Visa information om API-nyckel
    if not st.session_state.get("openai_api_key"):
        st.info("📝 **Ingen OpenAI API-nyckel angiven.** Transkribering kommer att fungera, men för att få en sammanfattning behöver du ange en API-nyckel i sidofältet.")
    else:
        st.success("✅ **OpenAI API-nyckel är konfigurerad.** Sammanfattningsfunktionen är aktiverad.")
    
    
    # SKAPA TVÅ FLIKAR FÖR FILUPPLADDNING OCH INSPELNING    
    tab1, tab2 = st.tabs(["Ladda upp fil", "Spela in ljud"])
    
    # FILUPPLADDARE
    with tab1:
        uploaded_file = st.file_uploader(
            "Ladda upp din ljud- eller videofil",
            type=["mp3", "wav", "flac", "mp4", "m4a", "aifc"],
            help="Max 2GB filstorlek",
        )
        
        if uploaded_file:
            # Kontrollerar om uppladdad fil redan har transkriberats
            current_file_hash = compute_file_hash(uploaded_file)
            
            # Om uppladdad fils hash är annorlunda än den i session state, återställ state
            if "file_hash" not in st.session_state or st.session_state.file_hash != current_file_hash:
                st.session_state.file_hash = current_file_hash
                
                if "transcribed" in st.session_state:
                    st.session_state.transcribed = None
                if "summarized" in st.session_state:
                    st.session_state.summarized = None
            
            # Om ljudet inte har transkriberats
            if st.session_state.transcribed is None:
                # Knapp för att starta bearbetning
                if st.button("Transkribera och sammanfatta"):
                    # Skickar ljud för konvertering till mp3 och komprimering
                    with st.spinner('Komprimerar ljudfil...'):
                        st.session_state.file_name_converted = convert_to_mono_and_compress(uploaded_file, uploaded_file.name, audio_dir)
                        if not st.session_state.file_name_converted:
                            st.error("Kunde inte bearbeta ljudfilen. Kontrollera filformatet och att ffmpeg är installerat.")
                            return
                        st.success('Ljudet komprimerat. Startar transkribering.')
                    
                    # Transkriberar ljud med KB Whisper
                    with st.spinner('Transkriberar... Detta kan ta en stund beroende på längden på inspelningen...'):
                        # Skapa progressionsbar
                        progress_bar = st.progress(0)
                        
                        # Skapa callback-funktion för att uppdatera progressionsbaren
                        def update_progress(progress_value):
                            progress_bar.progress(progress_value)
                        
                        # Anropa transkriberingsfunktionen med progress_callback
                        st.session_state.transcribed = transcribe_with_kb_whisper(
                            st.session_state.file_name_converted, 
                            uploaded_file.name, 
                            model_map_transcribe_model[st.session_state["transcribe_model"]],
                            "sv",
                            text_dir,
                            update_progress
                        )
                        
                        # Sätt progressionsbaren till 100% när den är klar
                        progress_bar.progress(1.0)
                        
                        st.success('Transkribering klar.')
                    
                    # Sammanfatta den transkriberade texten med OpenAI
                    with st.spinner('Sammanfattar transkriberingen...'):
                        st.session_state.summarized = summarize_text_openai(
                            st.session_state.transcribed, 
                            summary_length,
                            st.session_state.get("openai_api_key")
                        )
                        st.success('Sammanfattning klar.')
                        
                    st.balloons()
            
            # Om vi har transkriberad och sammanfattad text
            if st.session_state.transcribed is not None:
                # Skapar ett Word-dokument med den transkriberade texten
                document = Document()
                clean_text = st.session_state.transcribed.encode('utf-8', errors='replace').decode('utf-8')
                document.add_paragraph(clean_text)
                
                if st.session_state.summarized:
                    document.add_paragraph("\n\nSAMMANFATTNING:\n" + st.session_state.summarized)
                
                document.save(os.path.join(text_dir, uploaded_file.name + '.docx'))
                
                # Spara textfil
                with open(os.path.join(text_dir, uploaded_file.name + '.txt'), 'w', encoding='utf-8') as txt_file:
                    txt_file.write(clean_text)
                    if st.session_state.summarized:
                        txt_file.write("\n\nSAMMANFATTNING:\n" + st.session_state.summarized)
                
                with open(os.path.join(text_dir, uploaded_file.name + ".docx"), "rb") as docx_file:
                    docx_bytes = docx_file.read()
                
                # Skapar kolumner för nedladdningsknappar
                col1, col2 = st.columns(2)
                
                # Text nedladdning
                with col1:
                    with open(os.path.join(text_dir, uploaded_file.name + '.txt'), "rb") as file_txt:
                        st.download_button(
                            label = "Ladda ner som Text",
                            data = file_txt,
                            file_name = uploaded_file.name + '.txt',
                            mime = 'text/plain',
                        )
                
                # Word nedladdning
                with col2:
                    st.download_button(
                        label = "Ladda ner som Word",
                        data = docx_bytes,
                        file_name = uploaded_file.name + '.docx',
                        mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    )
                
                # Visa ljudet
                st.markdown("### Ljudfil")
                if st.session_state.file_name_converted is not None:
                    st.audio(st.session_state.file_name_converted, format='audio/wav')
                
                # Visa transkriberingen
                st.markdown("### Transkribering")
                st.write(st.session_state.transcribed)
                
                # Visa sammanfattningen
                if st.session_state.summarized:
                    st.markdown("### Sammanfattning")
                    st.write(st.session_state.summarized)
    
    # LJUDINSPELARE
    with tab2:
        audio = st.audio_input("Spela in ljud")
        
        if audio:
            # Öppna den sparade ljudfilen och beräkna dess hash
            current_file_hash = compute_file_hash(audio)
            
            # Om uppladdad fils hash är annorlunda än den i session state, återställ state
            if "file_hash" not in st.session_state or st.session_state.file_hash != current_file_hash:
                st.session_state.file_hash = current_file_hash
                
                if "transcribed" in st.session_state:
                    st.session_state.transcribed = None
                if "summarized" in st.session_state:
                    st.session_state.summarized = None
            
            # Om ljudet inte har transkriberats
            if st.session_state.transcribed is None:
                # Knapp för att starta bearbetning
                if st.button("Bearbeta inspelning"):
                    # Skapa en temporär fil för att spara inspelningen
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.webm') as tmp_file:
                        tmp_file.write(audio.getvalue())
                        tmp_path = tmp_file.name
                    
                    try:
                        audio_file = AudioSegment.from_file(tmp_path)
                        output_path = os.path.join(audio_dir, "inspelning.mp3")
                        audio_file.export(output_path, format="mp3", bitrate="16k")
                        os.unlink(tmp_path)  # Ta bort temporär fil
                    except Exception as e:
                        st.error(f"Kunde inte bearbeta inspelningen: {e}")
                        if os.path.exists(tmp_path):
                            os.unlink(tmp_path)
                        return
                    
                    # Transkribera
                    with st.spinner('Transkriberar... Detta kan ta en stund beroende på längden på inspelningen...'):
                        # Skapa progressionsbar
                        progress_bar = st.progress(0)
                        
                        # Skapa callback-funktion för att uppdatera progressionsbaren
                        def update_progress(progress_value):
                            progress_bar.progress(progress_value)
                        
                        # Anropa transkriberingsfunktionen med progress_callback
                        st.session_state.transcribed = transcribe_with_kb_whisper(
                            output_path, 
                            "inspelning.mp3", 
                            model_map_transcribe_model[st.session_state["transcribe_model"]],
                            "sv",
                            text_dir,
                            update_progress
                        )
                        
                        # Sätt progressionsbaren till 100% när den är klar
                        progress_bar.progress(1.0)
                        
                        st.success('Transkribering klar.')
                    
                    # Sammanfatta
                    with st.spinner('Sammanfattar transkriberingen...'):
                        st.session_state.summarized = summarize_text_openai(
                            st.session_state.transcribed, 
                            summary_length,
                            st.session_state.get("openai_api_key")
                        )
                        st.success('Sammanfattning klar.')
                    
                    st.balloons()
            
            # Om vi har transkriberad och sammanfattad text
            if st.session_state.transcribed is not None:
                recording_name = "inspelning.mp3"
                document = Document()
                clean_text = st.session_state.transcribed.encode('utf-8', errors='replace').decode('utf-8')
                document.add_paragraph(clean_text)
                
                if st.session_state.summarized:
                    document.add_paragraph("\n\nSAMMANFATTNING:\n" + st.session_state.summarized)
                
                document.save(os.path.join(text_dir, recording_name + '.docx'))
                
                # Spara textfil
                with open(os.path.join(text_dir, recording_name + '.txt'), 'w', encoding='utf-8') as txt_file:
                    txt_file.write(clean_text)
                    if st.session_state.summarized:
                        txt_file.write("\n\nSAMMANFATTNING:\n" + st.session_state.summarized)
                
                with open(os.path.join(text_dir, recording_name + '.docx'), "rb") as docx_file:
                    docx_bytes = docx_file.read()
                
                # Skapar kolumner för nedladdningsknappar
                col1, col2 = st.columns(2)
                
                # Text nedladdning
                with col1:
                    with open(os.path.join(text_dir, recording_name + '.txt'), "rb") as file_txt:
                        st.download_button(
                            label = "Ladda ner som Text",
                            data = file_txt,
                            file_name = recording_name + '.txt',
                            mime = 'text/plain',
                        )
                
                # Word nedladdning
                with col2:
                    st.download_button(
                        label = "Ladda ner som Word",
                        data = docx_bytes,
                        file_name = recording_name + '.docx',
                        mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    )
                
                # Visa transkriberingen
                st.markdown("### Transkribering")
                st.write(st.session_state.transcribed)
                
                # Visa sammanfattningen
                if st.session_state.summarized:
                    st.markdown("### Sammanfattning")
                    st.write(st.session_state.summarized)

if __name__ == "__main__":
    main()
